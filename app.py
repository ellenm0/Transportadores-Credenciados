import streamlit as st
import requests
import json
import base64
import uuid
import time
import unicodedata
from datetime import datetime, date
from io import BytesIO
import streamlit.components.v1 as components


st.set_page_config(
    page_title="Transportadores Licenciados e Credenciados",
    page_icon="🚛",
    layout="wide"
)


MODALIDADES = [
    "COLETA E TRANSPORTE DE RESÍDUOS NÃO PERIGOSOS",
    "COLETA DE RESÍDUOS PERIGOSOS",
    "COLETA DE RESÍDUOS VEGETAIS E DA CONSTRUÇÃO CIVIL COM FORNECIMENTO DE CAÇAMBA ESTACIONÁRIA",
    "COLETA DE RESÍDUOS VEGETAIS E DA CONSTRUÇÃO CIVIL PROVENIENTES DE ESCAVAÇÃO, DE DEMOLIÇÃO E DE SERVIÇOS DE TERRAPLENAGEM, POR MEIO DE CAÇAMBA BASCULANTE",
    "COLETA DE RESÍDUOS DE SERVIÇOS DE SAÚDE (HOSPITALAR E AMBULATORIAL)",
    "COLETA DE RESÍDUOS DE SERVIÇOS DE SAÚDE (AMBULATORIAL)",
    "COLETA DE RESÍDUOS RECICLÁVEIS",
    "COLETA DE PRODUTOS E EMBALAGENS OBJETOS DE LOGÍSTICA REVERSA",
    "COLETA DE EFLUENTES",
]


MODALIDADES_DESCRICOES = {
    "COLETA E TRANSPORTE DE RESÍDUOS NÃO PERIGOSOS": (
        "Destinada à coleta de resíduos sólidos caracterizados como da classe II, não "
        "perigosos, pela NBR 10.004, da Associação Brasileira de Normas Técnicas, "
        "gerados em atividades comerciais, industriais e de prestadores de serviços, "
        "em volume igual ou superior 100 (cem) litros por dia."
    ),
    "COLETA DE RESÍDUOS PERIGOSOS": (
        "Destinada à coleta de resíduos sólidos caracterizados como da classe I, "
        "perigosos, pela NBR 10.004, da Associação Brasileira de Normas Técnicas, "
        "gerados em atividades comerciais, industriais e de prestadores de serviços, "
        "independente do seu volume."
    ),
    "COLETA DE RESÍDUOS VEGETAIS E DA CONSTRUÇÃO CIVIL COM FORNECIMENTO DE CAÇAMBA ESTACIONÁRIA": (
        "Destinada, exclusivamente, à coleta de resíduos da construção civil "
        "provenientes de escavação, demolição e de serviços de terraplenagem, bem "
        "como dos provenientes de podas ou corte de árvores."
    ),
    "COLETA DE RESÍDUOS VEGETAIS E DA CONSTRUÇÃO CIVIL PROVENIENTES DE ESCAVAÇÃO, DE DEMOLIÇÃO E DE SERVIÇOS DE TERRAPLENAGEM, POR MEIO DE CAÇAMBA BASCULANTE": (
        "Destinada, exclusivamente, à coleta de resíduos da construção civil "
        "provenientes de escavação, demolição e de serviços de terraplenagem, bem "
        "como dos provenientes de podas ou corte de árvores."
    ),
    "COLETA DE RESÍDUOS DE SERVIÇOS DE SAÚDE (HOSPITALAR E AMBULATORIAL)": (
        "Destinada à coleta de resíduos de serviços de saúde gerados em "
        "estabelecimentos cujas atividades estejam relacionadas com a atenção à "
        "saúde humana ou animal, inclusive os serviços de assistência domiciliar, "
        "laboratórios analíticos de produtos para saúde, necrotérios, funerárias e "
        "serviços onde se realizem atividades de embalsamamento (tanatopraxia e "
        "somatoconservação), serviços de medicina legal, drogarias e farmácias, "
        "inclusive as de manipulação, estabelecimentos de ensino e pesquisa na área "
        "de saúde, centros de controle de zoonoses, distribuidores de produtos "
        "farmacêuticos, importadores, distribuidores de materiais e controles para "
        "diagnóstico in vitro, unidades móveis de atendimento à saúde, serviços de "
        "acupuntura, serviços de piercing e tatuagem, salões de beleza e estética, "
        "dentre outros afins."
    ),
    "COLETA DE RESÍDUOS DE SERVIÇOS DE SAÚDE (AMBULATORIAL)": (
        "Destinada, exclusivamente, à coleta de resíduos de serviços de saúde de "
        "pequenos geradores ou ambulatorial - assim definidos conforme NBR ABNT "
        "12.980/1993, cuja geração seja inferior a 700 L por semana ou a 150 L dia."
    ),
    "COLETA DE RESÍDUOS RECICLÁVEIS": (
        "Destinados à coleta de resíduos sólidos que devem retornar ao setor "
        "empresarial, após o uso pelo consumidor, para reaproveitamento, em seu "
        "ciclo ou em outros ciclos produtivos, ou outra destinação final "
        "ambientalmente adequada, na forma do art. 33 da Lei Federal nº 12.305, de "
        "02 de agosto de 2010."
    ),
    "COLETA DE PRODUTOS E EMBALAGENS OBJETOS DE LOGÍSTICA REVERSA": (
        "Destinados à coleta de resíduos sólidos que devem retornar ao setor "
        "empresarial, após o uso pelo consumidor, para reaproveitamento, em seu "
        "ciclo ou em outros ciclos produtivos, ou outra destinação final "
        "ambientalmente adequada, na forma do art. 33 da Lei Federal nº 12.305, de "
        "02 de agosto de 2010."
    ),
    "COLETA DE EFLUENTES": (
        "Destinados à coleta e transporte de efluentes domésticos e/ou industriais "
        "(limpa fossas) no município de Fortaleza, nos termos do Decreto Municipal "
        "nº 14.181, de 09 de março de 2018."
    ),
}


def agora_iso():
    return datetime.now().isoformat(timespec="seconds")


def normalizar_cnpj(valor):
    return "".join(
        c for c in str(valor or "")
        if c.isdigit()
    )


def formatar_cnpj(valor):
    numeros = normalizar_cnpj(valor)

    if len(numeros) == 14:
        return (
            f"{numeros[:2]}."
            f"{numeros[2:5]}."
            f"{numeros[5:8]}/"
            f"{numeros[8:12]}-"
            f"{numeros[12:]}"
        )

    return str(valor or "")


def remover_acentos(texto):
    """
    Remove acentos de um texto, mantendo as letras base, para permitir que
    a busca encontre "JOÃO" digitando "JOAO" (e vice-versa).
    """

    texto_normalizado = unicodedata.normalize(
        "NFKD",
        str(texto or "")
    )

    return "".join(
        caractere
        for caractere in texto_normalizado
        if not unicodedata.combining(caractere)
    )


def normalizar_busca(texto):
    return remover_acentos(texto).upper().strip()


def interpretar_validade(texto):
    """
    Tenta interpretar o campo de validade como uma data no formato
    brasileiro DD/MM/AAAA.

    Retorna um objeto `date` quando o texto é uma data válida, ou None
    quando não é (por exemplo "EM RENOVAÇÃO", "INDETERMINADA", vazio, ou
    qualquer outro texto livre). Isso garante que esses textos nunca sejam
    tratados como data vencida.
    """

    texto = str(texto or "").strip()

    if not texto:
        return None

    try:
        return datetime.strptime(texto, "%d/%m/%Y").date()
    except ValueError:
        return None


def gerar_status_transportador(transportador):
    """
    Calcula a situação atual do credenciamento SEM alterar o cadastro,
    comparando a validade com a data de hoje. A conta usa subtração de
    datas do Python (date - date), que já trata corretamente anos
    bissextos e qualquer diferença de calendário.

    Retorna um dicionário com:
      - efetivamente_ativo: se deve contar como ativo na listagem
      - vencido: se o credenciamento já passou da validade
      - rotulo: texto de status a ser exibido na tela
    """

    ativo_manual = transportador.get(
        "ativo",
        True
    )

    if not ativo_manual:

        return {
            "efetivamente_ativo": False,
            "vencido": False,
            "rotulo": "INATIVO"
        }

    data_validade = interpretar_validade(
        transportador.get(
            "validade_credenciamento",
            ""
        )
    )

    if data_validade is None:

        # Texto livre (EM RENOVAÇÃO, INDETERMINADA, vazio etc.):
        # não há data para calcular vencimento, então não é tratado
        # como vencido.
        return {
            "efetivamente_ativo": True,
            "vencido": False,
            "rotulo": "ATIVO"
        }

    hoje = date.today()

    dias_restantes = (
        data_validade - hoje
    ).days

    if dias_restantes < 0:

        return {
            "efetivamente_ativo": False,
            "vencido": True,
            "rotulo": "INATIVO — CREDENCIAMENTO VENCIDO"
        }

    if dias_restantes == 0:

        return {
            "efetivamente_ativo": True,
            "vencido": False,
            "rotulo": "ATIVO — vence hoje"
        }

    if dias_restantes <= 90:

        unidade = "dia" if dias_restantes == 1 else "dias"

        return {
            "efetivamente_ativo": True,
            "vencido": False,
            "rotulo": (
                f"ATIVO — {dias_restantes} {unidade} "
                f"para seu vencimento"
            )
        }

    return {
        "efetivamente_ativo": True,
        "vencido": False,
        "rotulo": "ATIVO"
    }


# ---------------------------------------------------------------------------
# Sessão HTTP reaproveitável (evita reabrir conexão a cada chamada ao
# GitHub, deixando o aplicativo mais rápido).
# ---------------------------------------------------------------------------

@st.cache_resource
def obter_sessao_http():
    return requests.Session()


def github_config():
    token = st.secrets.get(
        "GITHUB_TOKEN",
        ""
    )

    repo = st.secrets.get(
        "GITHUB_REPOSITORY",
        ""
    )

    branch = st.secrets.get(
        "GITHUB_BRANCH",
        "main"
    )

    data_dir = st.secrets.get(
        "GITHUB_DATA_DIR",
        "dados"
    )

    return (
        token,
        repo,
        branch,
        data_dir
    )


def github_request(
    method,
    url,
    token,
    **kwargs
):
    headers = kwargs.pop(
        "headers",
        {}
    )

    headers.update({
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json",
        "X-GitHub-Api-Version": "2022-11-28"
    })

    sessao = obter_sessao_http()

    return sessao.request(
        method,
        url,
        headers=headers,
        timeout=30,
        **kwargs
    )


def obter_sha_atual(
    nome_arquivo
):
    """
    Busca diretamente no GitHub o SHA mais atual
    do arquivo antes de fazer uma alteração.
    """

    token, repo, branch, data_dir = github_config()

    if not token or not repo:
        return (
            None,
            False,
            "Configure GITHUB_TOKEN e GITHUB_REPOSITORY nos Secrets."
        )

    data_dir = data_dir.strip("/")

    if data_dir:
        caminho = (
            f"{data_dir}/{nome_arquivo}"
        )
    else:
        caminho = nome_arquivo

    url = (
        f"https://api.github.com/repos/"
        f"{repo}/contents/{caminho}"
    )

    resposta = github_request(
        "GET",
        url,
        token,
        params={
            "ref": branch
        }
    )

    if resposta.status_code == 200:
        dados = resposta.json()

        return (
            dados.get("sha"),
            True,
            ""
        )

    if resposta.status_code == 404:
        # Arquivo ainda não existe.
        return (
            None,
            True,
            ""
        )

    return (
        None,
        False,
        (
            f"Erro ao consultar o arquivo no GitHub "
            f"({resposta.status_code}): "
            f"{resposta.text[:500]}"
        )
    )


def carregar_arquivo_github(
    nome_arquivo,
    padrao
):
    token, repo, branch, data_dir = github_config()

    if not token or not repo:
        return (
            padrao,
            None,
            False,
            "Configure GITHUB_TOKEN e GITHUB_REPOSITORY nos Secrets."
        )

    data_dir = data_dir.strip("/")

    if data_dir:
        caminho = (
            f"{data_dir}/{nome_arquivo}"
        )
    else:
        caminho = nome_arquivo

    url = (
        f"https://api.github.com/repos/"
        f"{repo}/contents/{caminho}"
    )

    resposta = github_request(
        "GET",
        url,
        token,
        params={
            "ref": branch
        }
    )

    if resposta.status_code == 200:

        dados = resposta.json()

        try:

            conteudo = base64.b64decode(
                dados["content"]
            ).decode("utf-8")

            return (
                json.loads(conteudo),
                dados.get("sha"),
                True,
                ""
            )

        except Exception as erro:

            return (
                padrao,
                None,
                False,
                (
                    f"Não foi possível ler "
                    f"{caminho}: {erro}"
                )
            )

    if resposta.status_code == 404:

        return (
            padrao,
            None,
            True,
            ""
        )

    return (
        padrao,
        None,
        False,
        (
            f"Erro ao acessar o GitHub "
            f"({resposta.status_code}): "
            f"{resposta.text[:500]}"
        )
    )


def salvar_arquivo_github(
    nome_arquivo,
    dados,
    sha=None,
    mensagem="Atualização dos dados"
):
    """
    Salva o arquivo no GitHub.

    IMPORTANTE:
    Antes de salvar, busca o SHA mais recente diretamente no GitHub. Isso
    evita conflitos de versão (SHA antigo) quando o arquivo tiver sido
    alterado entre o carregamento e o salvamento.

    Esta função NUNCA envia uma lista vazia "do nada": ela sempre recebe
    a lista completa já carregada (transportadores, historico ou
    relatorios), apenas com o item alterado/adicionado/removido — os
    demais registros são sempre preservados.
    """

    token, repo, branch, data_dir = github_config()

    if not token or not repo:
        return (
            False,
            "Configure GITHUB_TOKEN e GITHUB_REPOSITORY nos Secrets."
        )

    data_dir = data_dir.strip("/")

    if data_dir:
        caminho = (
            f"{data_dir}/{nome_arquivo}"
        )
    else:
        caminho = nome_arquivo

    url = (
        f"https://api.github.com/repos/"
        f"{repo}/contents/{caminho}"
    )

    sha_atual, consulta_ok, erro_consulta = obter_sha_atual(
        nome_arquivo
    )

    if not consulta_ok:

        return (
            False,
            erro_consulta
        )

    conteudo = json.dumps(
        dados,
        ensure_ascii=False,
        indent=2
    )

    payload = {
        "message": mensagem,
        "content": base64.b64encode(
            conteudo.encode("utf-8")
        ).decode("utf-8"),
        "branch": branch
    }

    # Se o arquivo já existe, usa o SHA mais recente.
    if sha_atual:
        payload["sha"] = sha_atual

    resposta = github_request(
        "PUT",
        url,
        token,
        json=payload
    )

    if resposta.status_code in (
        200,
        201
    ):
        return (
            True,
            ""
        )

    # ---------------------------------------------------------
    # Segunda tentativa:
    # caso o arquivo tenha sido alterado exatamente durante o
    # salvamento, busca novamente o SHA e tenta outra vez.
    # ---------------------------------------------------------

    if resposta.status_code in (
        409,
        422
    ):

        novo_sha, nova_consulta_ok, novo_erro = obter_sha_atual(
            nome_arquivo
        )

        if nova_consulta_ok:

            if novo_sha:
                payload["sha"] = novo_sha

            elif "sha" in payload:
                del payload["sha"]

            segunda_resposta = github_request(
                "PUT",
                url,
                token,
                json=payload
            )

            if segunda_resposta.status_code in (
                200,
                201
            ):
                return (
                    True,
                    ""
                )

            return (
                False,
                (
                    f"Erro ao salvar no GitHub "
                    f"({segunda_resposta.status_code}): "
                    f"{segunda_resposta.text[:500]}"
                )
            )

        return (
            False,
            novo_erro
        )

    return (
        False,
        (
            f"Erro ao salvar no GitHub "
            f"({resposta.status_code}): "
            f"{resposta.text[:500]}"
        )
    )


def carregar_dados():
    """
    Lê os três arquivos existentes no GitHub (transportadores.json,
    historico.json e relatorios.json) e retorna os dados EXATAMENTE como
    estão lá. Se um arquivo ainda não existir (404), retorna uma lista
    vazia apenas para esse arquivo específico — os arquivos que já
    existem nunca são tocados aqui, apenas lidos.
    """

    transportadores, sha_t, ok_t, erro_t = (
        carregar_arquivo_github(
            "transportadores.json",
            []
        )
    )

    historico, sha_h, ok_h, erro_h = (
        carregar_arquivo_github(
            "historico.json",
            []
        )
    )

    relatorios, sha_r, ok_r, erro_r = (
        carregar_arquivo_github(
            "relatorios.json",
            []
        )
    )

    if not ok_t:
        st.error(erro_t)

    if not ok_h:
        st.error(erro_h)

    if not ok_r:
        st.error(erro_r)

    if not isinstance(
        transportadores,
        list
    ):
        transportadores = []

    if not isinstance(
        historico,
        list
    ):
        historico = []

    if not isinstance(
        relatorios,
        list
    ):
        relatorios = []

    return (
        transportadores,
        historico,
        relatorios,
        sha_t,
        sha_h,
        sha_r
    )


def garantir_dados_no_github(
    transportadores,
    historico,
    relatorios,
    sha_t,
    sha_h,
    sha_r
):
    """
    Só cria um arquivo no GitHub quando ele realmente NÃO existe
    (sha is None, ou seja, a consulta retornou 404). Se o arquivo já
    existe, esta função não faz nada — nunca sobrescreve dados
    existentes com uma lista vazia.
    """

    token, repo, branch, data_dir = github_config()

    if not token or not repo:
        return

    if sha_t is None:

        salvar_arquivo_github(
            "transportadores.json",
            transportadores,
            None,
            "Criar cadastro inicial"
        )

    if sha_h is None:

        salvar_arquivo_github(
            "historico.json",
            historico,
            None,
            "Criar histórico inicial"
        )

    if sha_r is None:

        salvar_arquivo_github(
            "relatorios.json",
            relatorios,
            None,
            "Criar histórico de relatórios"
        )


# ---------------------------------------------------------------------------
# CACHE EM MEMÓRIA (por sessão do navegador)
#
# Sem isso, toda troca de aba/clique fazia 3 requisições ao GitHub de
# novo. Agora os dados são buscados uma vez e guardados em
# st.session_state; as telas usam essa cópia em memória. Como as listas
# (transportadores, historico, relatorios) são sempre alteradas "no
# lugar" (append, [:] = ..., etc.), qualquer alteração feita em uma tela
# aparece imediatamente nas outras, sem precisar buscar o GitHub de novo
# — e sem risco de mostrar dado desatualizado depois de salvar.
#
# O cache também expira sozinho depois de alguns minutos, e existe um
# botão manual "Atualizar dados" para forçar a releitura do GitHub a
# qualquer momento.
# ---------------------------------------------------------------------------

CACHE_TTL_SEGUNDOS = 300  # 5 minutos


def carregar_dados_cache(forcar=False):

    cache_expirado = False

    if "cache_timestamp" in st.session_state:
        cache_expirado = (
            time.time() - st.session_state["cache_timestamp"]
        ) > CACHE_TTL_SEGUNDOS

    precisa_buscar = (
        forcar
        or "cache_carregado" not in st.session_state
        or cache_expirado
    )

    if precisa_buscar:

        with st.spinner("Carregando dados do GitHub..."):

            transportadores, historico, relatorios, sha_t, sha_h, sha_r = (
                carregar_dados()
            )

            garantir_dados_no_github(
                transportadores,
                historico,
                relatorios,
                sha_t,
                sha_h,
                sha_r
            )

        st.session_state["transportadores"] = transportadores
        st.session_state["historico"] = historico
        st.session_state["relatorios"] = relatorios
        st.session_state["sha_t"] = sha_t
        st.session_state["sha_h"] = sha_h
        st.session_state["sha_r"] = sha_r
        st.session_state["cache_carregado"] = True
        st.session_state["cache_timestamp"] = time.time()

    return (
        st.session_state["transportadores"],
        st.session_state["historico"],
        st.session_state["relatorios"],
        st.session_state["sha_t"],
        st.session_state["sha_h"],
        st.session_state["sha_r"]
    )


def snapshot_transportador(
    transportador,
    acao
):
    return {
        "data": agora_iso(),

        "acao": acao,

        "transportador": json.loads(
            json.dumps(
                transportador,
                ensure_ascii=False
            )
        )
    }


def montar_endereco(
    transportador
):
    partes = []

    logradouro = transportador.get(
        "logradouro",
        ""
    ).strip()

    numero = transportador.get(
        "numero",
        ""
    ).strip()

    complemento = transportador.get(
        "complemento",
        ""
    ).strip()

    bairro = transportador.get(
        "bairro",
        ""
    ).strip()

    municipio = transportador.get(
        "municipio",
        ""
    ).strip()

    uf = transportador.get(
        "uf",
        ""
    ).strip()

    if logradouro:

        endereco = logradouro

        if numero:
            endereco += (
                f", Nº {numero}"
            )

        if complemento:
            endereco += (
                f", {complemento}"
            )

        partes.append(endereco)

    if bairro:
        partes.append(bairro)

    cidade_uf = " - ".join(
        x
        for x in [
            municipio,
            uf
        ]
        if x
    )

    if cidade_uf:
        partes.append(
            cidade_uf
        )

    return ", ".join(partes)


def montar_contato(
    transportador
):
    partes = []

    telefone = transportador.get(
        "telefone",
        ""
    ).strip()

    email = transportador.get(
        "email",
        ""
    ).strip()

    if telefone:
        partes.append(
            f"Fone: {telefone}"
        )

    if email:
        partes.append(
            f"E-mail: {email}"
        )

    return " ".join(partes)


def gerar_bloco_transportador(
    transportador
):
    linhas = []

    nome = transportador.get(
        "nome",
        ""
    ).strip()

    if nome:
        linhas.append(nome)

    cnpj = formatar_cnpj(
        transportador.get(
            "cnpj",
            ""
        )
    )

    if cnpj:
        linhas.append(
            f"CNPJ: {cnpj}"
        )

    endereco = montar_endereco(
        transportador
    )

    if endereco:
        linhas.append(endereco)

    contato = montar_contato(
        transportador
    )

    if contato:
        linhas.append(contato)

    tipo = transportador.get(
        "tipo_credenciamento",
        ""
    ).strip()

    numero = transportador.get(
        "numero_credenciamento",
        ""
    ).strip()

    validade_credenciamento = (
        transportador.get(
            "validade_credenciamento",
            ""
        ).strip()
    )

    if tipo and numero:

        linhas.append(
            f"{tipo}: {numero}"
        )

    elif numero:

        linhas.append(
            f"Credenciamento: {numero}"
        )

    if validade_credenciamento:

        linhas.append(
            f"VALIDADE: "
            f"{validade_credenciamento}"
        )

    descricao_licenca = (
        transportador.get(
            "descricao_licenca",
            ""
        ).strip()
    )

    validade_licenca = (
        transportador.get(
            "validade_licenca",
            ""
        ).strip()
    )

    if descricao_licenca:
        linhas.append(
            descricao_licenca
        )

    if validade_licenca:

        linhas.append(
            f"VALIDADE: "
            f"{validade_licenca}"
        )

    return "\n".join(linhas)


def gerar_relatorio(
    transportadores,
    data_atualizacao
):
    """
    Gera o texto completo do relatório, no formato oficial: título, data,
    e para cada modalidade o cabeçalho "MODALIDADE: ..." seguido do
    enunciado oficial da modalidade e das empresas credenciadas para ela.
    """

    linhas = [
        "RELAÇÃO DE TRANSPORTADORES LICENCIADOS E CREDENCIADOS",
        f"ATUALIZADA EM {data_atualizacao}",
        ""
    ]

    for modalidade in MODALIDADES:

        empresas = []

        for transportador in transportadores:

            if not transportador.get(
                "ativo",
                True
            ):
                continue

            modalidades = transportador.get(
                "modalidades",
                []
            )

            if modalidade in modalidades:
                empresas.append(
                    transportador
                )

        if not empresas:
            continue

        linhas.append(
            f"MODALIDADE: {modalidade}"
        )

        descricao_modalidade = MODALIDADES_DESCRICOES.get(
            modalidade,
            ""
        )

        if descricao_modalidade:
            linhas.append(
                descricao_modalidade
            )

        linhas.append("")

        empresas.sort(
            key=lambda x: x.get(
                "nome",
                ""
            ).upper()
        )

        vistos = set()

        for transportador in empresas:

            chave = normalizar_cnpj(
                transportador.get(
                    "cnpj",
                    ""
                )
            )

            if not chave:

                chave = transportador.get(
                    "id",
                    ""
                )

            if chave in vistos:
                continue

            vistos.add(chave)

            linhas.append(
                gerar_bloco_transportador(
                    transportador
                )
            )

            linhas.append("")

    while linhas and linhas[-1] == "":
        linhas.pop()

    return "\n".join(linhas)


def criar_docx(
    texto
):
    """
    Converte o texto do relatório em um .docx formatado:
      - Título e data: negrito, centralizados
      - "MODALIDADE: ...": negrito, centralizado
      - Enunciado da modalidade (texto logo abaixo do MODALIDADE):
        itálico, justificado
      - Nome da empresa (linha imediatamente antes de "CNPJ: ..."):
        negrito e sublinhado
      - Demais linhas: texto normal
      - Espaçamento entre linhas de aproximadamente 1,15 em todo o
        documento
    """

    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    documento = Document()

    secao = documento.sections[0]

    secao.top_margin = Pt(50)
    secao.bottom_margin = Pt(50)
    secao.left_margin = Pt(60)
    secao.right_margin = Pt(60)

    linhas = texto.splitlines()

    for indice, linha in enumerate(linhas):

        linha_limpa = linha.strip()

        proxima_nao_vazia = ""

        for seguinte in linhas[indice + 1:]:
            if seguinte.strip():
                proxima_nao_vazia = seguinte.strip()
                break

        anterior_nao_vazia = ""

        for anterior in reversed(linhas[:indice]):
            if anterior.strip():
                anterior_nao_vazia = anterior.strip()
                break

        paragrafo = documento.add_paragraph()

        formato = paragrafo.paragraph_format
        formato.space_after = Pt(4)
        formato.line_spacing = 1.15

        if linha_limpa.startswith(
            "RELAÇÃO DE TRANSPORTADORES"
        ):

            formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
            formato.space_after = Pt(6)

            run = paragrafo.add_run(linha_limpa)
            run.bold = True
            run.font.size = Pt(15)

        elif linha_limpa.startswith(
            "ATUALIZADA EM"
        ):

            formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
            formato.space_after = Pt(14)

            run = paragrafo.add_run(linha_limpa)
            run.bold = True
            run.font.size = Pt(12)

        elif linha_limpa.startswith(
            "MODALIDADE:"
        ):

            formato.alignment = WD_ALIGN_PARAGRAPH.CENTER
            formato.space_before = Pt(16)
            formato.space_after = Pt(4)

            run = paragrafo.add_run(linha_limpa)
            run.bold = True
            run.font.size = Pt(12)

        elif (
            linha_limpa
            and proxima_nao_vazia.startswith("CNPJ:")
        ):

            # Nome da empresa: negrito + sublinhado
            formato.space_before = Pt(8)

            run = paragrafo.add_run(linha_limpa)
            run.bold = True
            run.underline = True
            run.font.size = Pt(11)

        elif (
            anterior_nao_vazia.startswith("MODALIDADE:")
            and linha_limpa
        ):

            # Enunciado/descrição oficial da modalidade
            formato.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            formato.space_after = Pt(10)

            run = paragrafo.add_run(linha_limpa)
            run.italic = True
            run.font.size = Pt(10)

        elif linha_limpa:

            run = paragrafo.add_run(linha_limpa)
            run.font.size = Pt(10)

        # Linhas em branco viram parágrafos vazios (apenas espaçamento).

    arquivo = BytesIO()

    documento.save(
        arquivo
    )

    arquivo.seek(0)

    return arquivo


def copiar_texto_componente(
    texto,
    identificador
):

    texto_js = json.dumps(
        texto,
        ensure_ascii=False
    )

    html = f"""
    <div style="
        display:flex;
        align-items:center;
        gap:10px;
        font-family:Arial,sans-serif;
        margin:4px 0 8px 0;
    ">

        <button
            id="btn_{identificador}"
            style="
                background:#ff4b4b;
                color:white;
                border:none;
                border-radius:6px;
                padding:10px 18px;
                cursor:pointer;
                font-size:14px;
                font-weight:600;
            "
        >
            Copiar relatório
        </button>

        <span
            id="msg_{identificador}"
            style="
                font-size:14px;
                color:#228b22;
            "
        ></span>

    </div>

    <script>

        const texto = {texto_js};

        document
            .getElementById(
                "btn_{identificador}"
            )
            .addEventListener(
                "click",
                async function() {{

                    try {{

                        await navigator
                            .clipboard
                            .writeText(texto);

                        document
                            .getElementById(
                                "msg_{identificador}"
                            )
                            .innerText =
                            "Relatório copiado!";

                    }} catch (erro) {{

                        const area =
                            document.createElement(
                                "textarea"
                            );

                        area.value = texto;

                        document.body.appendChild(
                            area
                        );

                        area.select();

                        document.execCommand(
                            "copy"
                        );

                        document.body.removeChild(
                            area
                        );

                        document
                            .getElementById(
                                "msg_{identificador}"
                            )
                            .innerText =
                            "Relatório copiado!";
                    }}

                }}
            );

    </script>
    """

    components.html(
        html,
        height=55
    )


def validar_transportador(
    transportador,
    transportadores,
    id_atual=None
):

    erros = []

    if not transportador[
        "nome"
    ].strip():

        erros.append(
            "Informe o Nome/Razão Social."
        )

    cnpj = normalizar_cnpj(
        transportador["cnpj"]
    )

    if len(cnpj) != 14:

        erros.append(
            "Informe um CNPJ com 14 dígitos."
        )

    if not transportador[
        "logradouro"
    ].strip():

        erros.append(
            "Informe a Rua/Logradouro."
        )

    if not transportador[
        "numero"
    ].strip():

        erros.append(
            "Informe o Número."
        )

    if not transportador[
        "bairro"
    ].strip():

        erros.append(
            "Informe o Bairro."
        )

    if not transportador[
        "municipio"
    ].strip():

        erros.append(
            "Informe o Município."
        )

    if not transportador[
        "uf"
    ].strip():

        erros.append(
            "Informe a UF."
        )

    if not transportador[
        "modalidades"
    ]:

        erros.append(
            "Selecione pelo menos uma modalidade."
        )

    for existente in transportadores:

        if existente.get(
            "id"
        ) == id_atual:

            continue

        cnpj_existente = normalizar_cnpj(
            existente.get(
                "cnpj",
                ""
            )
        )

        if cnpj_existente == cnpj:

            erros.append(
                "Já existe um transportador cadastrado "
                "com este CNPJ."
            )

            break

    transportador[
        "modalidades"
    ] = list(
        dict.fromkeys(
            transportador[
                "modalidades"
            ]
        )
    )

    return erros


def carregar_formulario(
    transportador
):

    st.session_state[
        "form_nome"
    ] = transportador.get(
        "nome",
        ""
    )

    st.session_state[
        "form_cnpj"
    ] = transportador.get(
        "cnpj",
        ""
    )

    st.session_state[
        "form_logradouro"
    ] = transportador.get(
        "logradouro",
        ""
    )

    st.session_state[
        "form_numero"
    ] = transportador.get(
        "numero",
        ""
    )

    st.session_state[
        "form_complemento"
    ] = transportador.get(
        "complemento",
        ""
    )

    st.session_state[
        "form_bairro"
    ] = transportador.get(
        "bairro",
        ""
    )

    st.session_state[
        "form_municipio"
    ] = transportador.get(
        "municipio",
        ""
    )

    st.session_state[
        "form_uf"
    ] = transportador.get(
        "uf",
        "CE"
    )

    st.session_state[
        "form_telefone"
    ] = transportador.get(
        "telefone",
        ""
    )

    st.session_state[
        "form_email"
    ] = transportador.get(
        "email",
        ""
    )

    st.session_state[
        "form_tipo_credenciamento"
    ] = transportador.get(
        "tipo_credenciamento",
        "Processo de Credenciamento"
    )

    st.session_state[
        "form_numero_credenciamento"
    ] = transportador.get(
        "numero_credenciamento",
        ""
    )

    st.session_state[
        "form_validade_credenciamento"
    ] = transportador.get(
        "validade_credenciamento",
        ""
    )

    st.session_state[
        "form_descricao_licenca"
    ] = transportador.get(
        "descricao_licenca",
        ""
    )

    st.session_state[
        "form_validade_licenca"
    ] = transportador.get(
        "validade_licenca",
        ""
    )

    st.session_state[
        "form_modalidades"
    ] = transportador.get(
        "modalidades",
        []
    )

    st.session_state[
        "editando_id"
    ] = transportador.get(
        "id",
        ""
    )


def limpar_formulario():

    chaves = [
        "form_nome",
        "form_cnpj",
        "form_logradouro",
        "form_numero",
        "form_complemento",
        "form_bairro",
        "form_municipio",
        "form_uf",
        "form_telefone",
        "form_email",
        "form_numero_credenciamento",
        "form_validade_credenciamento",
        "form_descricao_licenca",
        "form_validade_licenca",
        "form_modalidades"
    ]

    for chave in chaves:

        st.session_state.pop(
            chave,
            None
        )

    st.session_state[
        "form_tipo_credenciamento"
    ] = "Processo de Credenciamento"

    st.session_state[
        "editando_id"
    ] = None


def solicitar_pagina(pagina):
    """
    Agenda a troca de página para a PRÓXIMA execução do script, em vez de
    alterar st.session_state["pagina"] diretamente.

    Isso é o que evita o StreamlitWidgetAlreadyInstantiatedError: o menu
    lateral (st.sidebar.radio) é criado com key="pagina" logo no início
    de main(). Se qualquer botão, mais abaixo na mesma execução, tentasse
    escrever em st.session_state["pagina"] depois que esse widget já foi
    criado, o Streamlit recusa a alteração e gera esse erro.

    Por isso, os botões apenas registram o pedido em
    st.session_state["pagina_solicitada"] (uma chave comum, sem nenhum
    widget associado) e chamam st.rerun(). No início da PRÓXIMA
    execução, main() lê esse pedido e só então atualiza
    st.session_state["pagina"] — sempre ANTES do menu lateral ser criado.
    """

    st.session_state["pagina_solicitada"] = pagina


def tela_formulario(
    transportadores,
    historico,
    sha_t,
    sha_h
):

    editando_id = st.session_state.get(
        "editando_id"
    )

    existente = None

    for transportador in transportadores:

        if transportador.get(
            "id"
        ) == editando_id:

            existente = transportador

            break

    if existente:

        st.title(
            "Editar Transportador"
        )

        st.info(
            f"Editando: "
            f"{existente.get('nome', '')} "
            f"— CNPJ "
            f"{formatar_cnpj(existente.get('cnpj', ''))}"
        )

    else:

        st.title(
            "Novo Transportador"
        )

    with st.form(
        "form_transportador",
        clear_on_submit=False
    ):

        st.subheader(
            "Dados cadastrais"
        )

        coluna1, coluna2 = st.columns(2)

        with coluna1:

            nome = st.text_input(
                "Nome/Razão Social *",
                value=st.session_state.get(
                    "form_nome",
                    existente.get(
                        "nome",
                        ""
                    )
                    if existente
                    else ""
                )
            )

            cnpj = st.text_input(
                "CNPJ *",
                value=st.session_state.get(
                    "form_cnpj",
                    existente.get(
                        "cnpj",
                        ""
                    )
                    if existente
                    else ""
                )
            )

            logradouro = st.text_input(
                "Rua/Logradouro *",
                value=st.session_state.get(
                    "form_logradouro",
                    existente.get(
                        "logradouro",
                        ""
                    )
                    if existente
                    else ""
                )
            )

            numero = st.text_input(
                "Número *",
                value=st.session_state.get(
                    "form_numero",
                    existente.get(
                        "numero",
                        ""
                    )
                    if existente
                    else ""
                )
            )

            complemento = st.text_input(
                "Complemento",
                value=st.session_state.get(
                    "form_complemento",
                    existente.get(
                        "complemento",
                        ""
                    )
                    if existente
                    else ""
                )
            )

        with coluna2:

            bairro = st.text_input(
                "Bairro *",
                value=st.session_state.get(
                    "form_bairro",
                    existente.get(
                        "bairro",
                        ""
                    )
                    if existente
                    else ""
                )
            )

            municipio = st.text_input(
                "Município *",
                value=st.session_state.get(
                    "form_municipio",
                    existente.get(
                        "municipio",
                        ""
                    )
                    if existente
                    else ""
                )
            )

            uf = st.text_input(
                "UF *",
                value=st.session_state.get(
                    "form_uf",
                    existente.get(
                        "uf",
                        "CE"
                    )
                    if existente
                    else "CE"
                )
            )

            telefone = st.text_input(
                "Telefone",
                value=st.session_state.get(
                    "form_telefone",
                    existente.get(
                        "telefone",
                        ""
                    )
                    if existente
                    else ""
                )
            )

            email = st.text_input(
                "E-mail",
                value=st.session_state.get(
                    "form_email",
                    existente.get(
                        "email",
                        ""
                    )
                    if existente
                    else ""
                )
            )

        st.subheader(
            "Credenciamento"
        )

        tipo_atual = st.session_state.get(
            "form_tipo_credenciamento",
            existente.get(
                "tipo_credenciamento",
                "Processo de Credenciamento"
            )
            if existente
            else "Processo de Credenciamento"
        )

        opcoes_credenciamento = [
            "Processo de Credenciamento",
            "Número de Credenciamento"
        ]

        if tipo_atual not in opcoes_credenciamento:

            tipo_atual = (
                opcoes_credenciamento[0]
            )

        tipo_credenciamento = st.selectbox(
            "Tipo",
            opcoes_credenciamento,
            index=opcoes_credenciamento.index(
                tipo_atual
            )
        )

        numero_credenciamento = st.text_input(
            "Número do processo/credenciamento",
            value=st.session_state.get(
                "form_numero_credenciamento",
                existente.get(
                    "numero_credenciamento",
                    ""
                )
                if existente
                else ""
            )
        )

        validade_credenciamento = st.text_input(
            "Validade do credenciamento",
            value=st.session_state.get(
                "form_validade_credenciamento",
                existente.get(
                    "validade_credenciamento",
                    ""
                )
                if existente
                else ""
            ),
            placeholder=(
                "Ex.: 24/05/2026, "
                "EM RENOVAÇÃO ou INDETERMINADA"
            ),
            help=(
                "Se você digitar uma data no formato DD/MM/AAAA, o "
                "sistema calcula automaticamente quando o credenciamento "
                "vence. Textos como \"EM RENOVAÇÃO\" ou \"INDETERMINADA\" "
                "nunca são tratados como vencidos."
            )
        )

        st.subheader(
            "Licenciamento"
        )

        descricao_licenca = st.text_area(
            "Descrição da licença",
            value=st.session_state.get(
                "form_descricao_licenca",
                existente.get(
                    "descricao_licenca",
                    ""
                )
                if existente
                else ""
            ),
            height=100,
            placeholder=(
                "Digite livremente a descrição da licença."
            )
        )

        validade_licenca = st.text_input(
            "Validade da licença",
            value=st.session_state.get(
                "form_validade_licenca",
                existente.get(
                    "validade_licenca",
                    ""
                )
                if existente
                else ""
            ),
            placeholder=(
                "Ex.: 20/12/2027, "
                "EM RENOVAÇÃO ou INDETERMINADA"
            )
        )

        st.subheader(
            "Modalidades"
        )

        modalidades = st.multiselect(
            "Selecione uma ou várias modalidades *",
            MODALIDADES,
            default=st.session_state.get(
                "form_modalidades",
                existente.get(
                    "modalidades",
                    []
                )
                if existente
                else []
            )
        )

        coluna_salvar, coluna_cancelar = st.columns(2)

        with coluna_salvar:

            salvar = st.form_submit_button(
                "Salvar transportador",
                type="primary",
                use_container_width=True
            )

        with coluna_cancelar:

            cancelar = st.form_submit_button(
                "Cancelar",
                use_container_width=True
            )

    if cancelar:

        limpar_formulario()

        solicitar_pagina(
            "Transportadores"
        )

        st.rerun()

    if salvar:

        novo = {

            "id": (
                existente.get("id")
                if existente
                else str(uuid.uuid4())
            ),

            "nome": nome.strip(),

            "cnpj": normalizar_cnpj(
                cnpj
            ),

            "logradouro": (
                logradouro.strip()
            ),

            "numero": (
                numero.strip()
            ),

            "complemento": (
                complemento.strip()
            ),

            "bairro": (
                bairro.strip()
            ),

            "municipio": (
                municipio.strip()
            ),

            "uf": (
                uf.strip().upper()
            ),

            "telefone": (
                telefone.strip()
            ),

            "email": (
                email.strip()
            ),

            "tipo_credenciamento": (
                tipo_credenciamento
            ),

            "numero_credenciamento": (
                numero_credenciamento.strip()
            ),

            "validade_credenciamento": (
                validade_credenciamento.strip()
            ),

            "descricao_licenca": (
                descricao_licenca.strip()
            ),

            "validade_licenca": (
                validade_licenca.strip()
            ),

            "modalidades": list(
                dict.fromkeys(
                    modalidades
                )
            ),

            "ativo": (
                existente.get(
                    "ativo",
                    True
                )
                if existente
                else True
            ),

            "criado_em": (
                existente.get(
                    "criado_em",
                    agora_iso()
                )
                if existente
                else agora_iso()
            ),

            "atualizado_em": agora_iso()
        }

        erros = validar_transportador(
            novo,
            transportadores,
            existente.get("id")
            if existente
            else None
        )

        if erros:

            for erro in erros:
                st.error(erro)

        else:

            if existente:

                historico.append(
                    snapshot_transportador(
                        existente,
                        "ANTES DA EDIÇÃO"
                    )
                )

                transportadores[:] = [
                    novo
                    if x.get("id")
                    == existente.get("id")
                    else x
                    for x in transportadores
                ]

                mensagem = (
                    f"Editar transportador "
                    f"{novo['nome']}"
                )

            else:

                transportadores.append(
                    novo
                )

                mensagem = (
                    f"Cadastrar transportador "
                    f"{novo['nome']}"
                )

            # -------------------------------------------------
            # SALVA O CADASTRO
            # transportadores já contém TODOS os registros
            # anteriores + a alteração atual — nada é apagado.
            # -------------------------------------------------

            ok1, erro1 = salvar_arquivo_github(
                "transportadores.json",
                transportadores,
                sha_t,
                mensagem
            )

            if not ok1:

                st.error(
                    erro1
                )

            else:

                # ---------------------------------------------
                # SALVA O HISTÓRICO (também preservando tudo
                # que já existia).
                # ---------------------------------------------

                ok2, erro2 = salvar_arquivo_github(
                    "historico.json",
                    historico,
                    sha_h,
                    (
                        f"Registrar histórico - "
                        f"{novo['nome']}"
                    )
                )

                if not ok2:

                    st.warning(
                        "Cadastro salvo, mas o histórico "
                        "não pôde ser atualizado: "
                        f"{erro2}"
                    )

                st.success(
                    "Transportador salvo com sucesso."
                )

                limpar_formulario()

                # Como a lista "transportadores" em session_state foi
                # alterada no lugar (é o mesmo objeto), o cache já fica
                # atualizado automaticamente — não é preciso recarregar
                # do GitHub para ver o resultado.

                solicitar_pagina(
                    "Transportadores"
                )

                st.rerun()


def tela_transportadores(
    transportadores,
    historico,
    sha_t,
    sha_h
):

    st.title(
        "Transportadores Cadastrados"
    )

    busca = st.text_input(
        "Pesquisar por Nome ou CNPJ",
        placeholder=(
            "Digite parte do nome ou CNPJ"
        )
    )

    mostrar_inativos = st.checkbox(
        "Mostrar também os inativos",
        value=False,
        help=(
            "Também exibe transportadores desativados manualmente e "
            "transportadores com credenciamento vencido."
        )
    )

    # Busca sem diferenciar acentuação nem maiúsculas/minúsculas:
    # "joao", "JOÃO" e "João" agora encontram o mesmo cadastro. O CNPJ
    # continua sendo comparado tanto formatado quanto só com números.
    termo = normalizar_busca(busca)

    filtrados = []

    for transportador in transportadores:

        status_info = gerar_status_transportador(
            transportador
        )

        if (
            not mostrar_inativos
            and not status_info["efetivamente_ativo"]
        ):
            continue

        texto = normalizar_busca(
            f"{transportador.get('nome', '')} "
            f"{formatar_cnpj(transportador.get('cnpj', ''))} "
            f"{normalizar_cnpj(transportador.get('cnpj', ''))}"
        )

        if (
            not termo
            or termo in texto
        ):

            filtrados.append(
                (transportador, status_info)
            )

    filtrados.sort(
        key=lambda par: par[0].get(
            "nome",
            ""
        ).upper()
    )

    st.caption(
        f"{len(filtrados)} "
        f"transportador(es) encontrado(s)."
    )

    if not filtrados:

        st.info(
            "Nenhum transportador encontrado."
        )

        return

    for transportador, status_info in filtrados:

        with st.container(
            border=True
        ):

            coluna_info, coluna_acoes = st.columns(
                [5, 2]
            )

            with coluna_info:

                st.markdown(
                    f"### "
                    f"{transportador.get('nome', '')}"
                )

                st.write(
                    f"**CNPJ:** "
                    f"{formatar_cnpj(transportador.get('cnpj', ''))}"
                )

                st.write(
                    f"**Município:** "
                    f"{transportador.get('municipio', '')}"
                )

                if status_info["vencido"]:

                    # Destaque discreto: só o rótulo de status recebe
                    # cor, o restante do cartão permanece neutro.
                    st.markdown(
                        "<div style='display:inline-block; "
                        "padding:3px 10px; border-radius:4px; "
                        "background-color:#fdecea; color:#b3261e; "
                        "font-weight:600; font-size:0.85em; "
                        "margin-top:2px;'>"
                        f"Status: {status_info['rotulo']}"
                        "</div>",
                        unsafe_allow_html=True
                    )

                else:

                    st.write(
                        f"**Status:** {status_info['rotulo']}"
                    )

                modalidades_texto = "; ".join(
                    transportador.get(
                        "modalidades",
                        []
                    )
                )

                st.write(
                    f"**Modalidades:** "
                    f"{modalidades_texto}"
                )

            with coluna_acoes:

                if st.button(
                    "Editar",
                    key=f"editar_{transportador['id']}",
                    use_container_width=True
                ):

                    carregar_formulario(
                        transportador
                    )

                    solicitar_pagina(
                        "Novo Transportador"
                    )

                    st.rerun()

                if transportador.get(
                    "ativo",
                    True
                ):

                    if st.button(
                        "Desativar",
                        key=f"desativar_{transportador['id']}",
                        use_container_width=True
                    ):

                        historico.append(
                            snapshot_transportador(
                                transportador,
                                "ANTES DA DESATIVAÇÃO"
                            )
                        )

                        transportador[
                            "ativo"
                        ] = False

                        transportador[
                            "atualizado_em"
                        ] = agora_iso()

                        ok1, erro1 = (
                            salvar_arquivo_github(
                                "transportadores.json",
                                transportadores,
                                sha_t,
                                (
                                    "Desativar transportador "
                                    f"{transportador['nome']}"
                                )
                            )
                        )

                        if ok1:

                            ok2, erro2 = (
                                salvar_arquivo_github(
                                    "historico.json",
                                    historico,
                                    sha_h,
                                    (
                                        "Registrar desativação - "
                                        f"{transportador['nome']}"
                                    )
                                )
                            )

                            if not ok2:

                                st.warning(
                                    "Transportador desativado, "
                                    "mas o histórico não pôde "
                                    "ser atualizado: "
                                    f"{erro2}"
                                )

                            st.success(
                                "Transportador desativado."
                            )

                            st.rerun()

                        else:

                            st.error(
                                erro1
                            )

                else:

                    if st.button(
                        "Reativar",
                        key=f"reativar_{transportador['id']}",
                        use_container_width=True
                    ):

                        historico.append(
                            snapshot_transportador(
                                transportador,
                                "ANTES DA REATIVAÇÃO"
                            )
                        )

                        transportador[
                            "ativo"
                        ] = True

                        transportador[
                            "atualizado_em"
                        ] = agora_iso()

                        ok1, erro1 = (
                            salvar_arquivo_github(
                                "transportadores.json",
                                transportadores,
                                sha_t,
                                (
                                    "Reativar transportador "
                                    f"{transportador['nome']}"
                                )
                            )
                        )

                        if ok1:

                            ok2, erro2 = (
                                salvar_arquivo_github(
                                    "historico.json",
                                    historico,
                                    sha_h,
                                    (
                                        "Registrar reativação - "
                                        f"{transportador['nome']}"
                                    )
                                )
                            )

                            if not ok2:

                                st.warning(
                                    "Transportador reativado, "
                                    "mas o histórico não pôde "
                                    "ser atualizado: "
                                    f"{erro2}"
                                )

                            st.success(
                                "Transportador reativado."
                            )

                            st.rerun()

                        else:

                            st.error(
                                erro1
                            )


def tela_relatorio(
    transportadores,
    relatorios,
    sha_r
):

    st.title(
        "Gerar Relatório"
    )

    data_padrao = st.session_state.get(
        "data_relatorio",
        date.today()
    )

    data_escolhida = st.date_input(
        "Data de atualização",
        value=data_padrao,
        format="DD/MM/YYYY"
    )

    data_formatada = (
        data_escolhida.strftime(
            "%d/%m/%Y"
        )
    )

    ativos = [
        t
        for t in transportadores
        if t.get(
            "ativo",
            True
        )
    ]

    st.caption(
        f"Serão considerados "
        f"{len(ativos)} "
        f"transportador(es) ativo(s)."
    )

    if st.button(
        "Gerar relatório",
        type="primary",
        use_container_width=True
    ):

        texto = gerar_relatorio(
            transportadores,
            data_formatada
        )

        registro = {
            "id": str(
                uuid.uuid4()
            ),

            "gerado_em": agora_iso(),

            "data_atualizacao": (
                data_formatada
            ),

            "texto": texto
        }

        relatorios.insert(
            0,
            registro
        )

        ok, erro = (
            salvar_arquivo_github(
                "relatorios.json",
                relatorios,
                sha_r,
                (
                    "Salvar relatório de "
                    f"{data_formatada}"
                )
            )
        )

        if ok:

            st.session_state[
                "relatorio_atual"
            ] = texto

            st.session_state[
                "relatorio_registro"
            ] = registro

            st.success(
                "Relatório gerado e salvo no histórico."
            )

        else:

            st.error(
                erro
            )

    texto = st.session_state.get(
        "relatorio_atual",
        ""
    )

    if texto:

        st.divider()

        st.subheader(
            "Relatório"
        )

        st.text_area(
            "Conteúdo completo",
            value=texto,
            height=600,
            key="area_relatorio"
        )

        copiar_texto_componente(
            texto,
            "relatorio_atual"
        )

        coluna1, coluna2 = st.columns(2)

        with coluna1:

            st.download_button(
                "Baixar TXT",
                data=texto.encode(
                    "utf-8"
                ),
                file_name=(
                    "relatorio_transportadores_"
                    f"{data_formatada.replace('/', '-')}.txt"
                ),
                mime="text/plain",
                use_container_width=True
            )

        with coluna2:

            documento = criar_docx(
                texto
            )

            st.download_button(
                "Baixar Word",
                data=documento,
                file_name=(
                    "relatorio_transportadores_"
                    f"{data_formatada.replace('/', '-')}.docx"
                ),
                mime=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
                use_container_width=True
            )


def tela_relatorios_anteriores(
    relatorios,
    sha_r
):

    st.title(
        "Relatórios Anteriores"
    )

    if not relatorios:

        st.info(
            "Nenhum relatório foi gerado ainda."
        )

        return

    for indice, relatorio in enumerate(
        relatorios
    ):

        # Garante que todo relatório tenha um id estável, mesmo os mais
        # antigos que por algum motivo não tenham um.
        if not relatorio.get("id"):
            relatorio["id"] = str(uuid.uuid4())

        relatorio_id = relatorio["id"]

        data_atualizacao = (
            relatorio.get(
                "data_atualizacao",
                ""
            )
        )

        gerado_em = (
            relatorio.get(
                "gerado_em",
                ""
            )
        )

        titulo = (
            f"Relatório de "
            f"{data_atualizacao}"
            f" — gerado em "
            f"{gerado_em}"
        )

        with st.expander(
            titulo
        ):

            texto = relatorio.get(
                "texto",
                ""
            )

            st.text_area(
                "Conteúdo",
                value=texto,
                height=400,
                key=f"relatorio_antigo_{relatorio_id}"
            )

            copiar_texto_componente(
                texto,
                f"relatorio_antigo_{indice}"
            )

            coluna1, coluna2, coluna3 = st.columns(3)

            with coluna1:

                st.download_button(
                    "Baixar TXT",
                    data=texto.encode(
                        "utf-8"
                    ),
                    file_name=(
                        "relatorio_transportadores_"
                        f"{str(data_atualizacao).replace('/', '-')}.txt"
                    ),
                    mime="text/plain",
                    key=f"download_txt_{relatorio_id}",
                    use_container_width=True
                )

            with coluna2:

                documento = criar_docx(
                    texto
                )

                st.download_button(
                    "Baixar Word",
                    data=documento,
                    file_name=(
                        "relatorio_transportadores_"
                        f"{str(data_atualizacao).replace('/', '-')}.docx"
                    ),
                    mime=(
                        "application/vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                    key=f"download_docx_{relatorio_id}",
                    use_container_width=True
                )

            with coluna3:

                chave_confirmacao = (
                    f"confirmar_exclusao_{relatorio_id}"
                )

                if not st.session_state.get(
                    chave_confirmacao,
                    False
                ):

                    if st.button(
                        "🗑️ Excluir",
                        key=f"excluir_{relatorio_id}",
                        use_container_width=True
                    ):

                        st.session_state[
                            chave_confirmacao
                        ] = True

                        st.rerun()

                else:

                    st.warning(
                        "Tem certeza que deseja excluir o relatório "
                        f"de {data_atualizacao}? Esta ação não pode "
                        "ser desfeita."
                    )

                    coluna_sim, coluna_nao = st.columns(2)

                    with coluna_sim:

                        if st.button(
                            "Sim, excluir",
                            key=f"confirmar_sim_{relatorio_id}",
                            type="primary",
                            use_container_width=True
                        ):

                            relatorios[:] = [
                                r
                                for r in relatorios
                                if r.get("id") != relatorio_id
                            ]

                            ok, erro = salvar_arquivo_github(
                                "relatorios.json",
                                relatorios,
                                sha_r,
                                (
                                    "Excluir relatório de "
                                    f"{data_atualizacao}"
                                )
                            )

                            st.session_state.pop(
                                chave_confirmacao,
                                None
                            )

                            if ok:
                                st.success(
                                    "Relatório excluído."
                                )
                            else:
                                st.error(
                                    erro
                                )

                            st.rerun()

                    with coluna_nao:

                        if st.button(
                            "Cancelar",
                            key=f"confirmar_nao_{relatorio_id}",
                            use_container_width=True
                        ):

                            st.session_state.pop(
                                chave_confirmacao,
                                None
                            )

                            st.rerun()


def tela_inicio(
    transportadores,
    relatorios
):

    st.title(
        "Transportadores Licenciados e Credenciados"
    )

    st.caption(
        "Cadastro, gerenciamento e geração da relação por modalidade."
    )

    ativos = sum(
        1
        for t in transportadores
        if t.get(
            "ativo",
            True
        )
    )

    inativos = (
        len(transportadores)
        - ativos
    )

    modalidades_ativas = sum(
        len(
            set(
                t.get(
                    "modalidades",
                    []
                )
            )
        )
        for t in transportadores
        if t.get(
            "ativo",
            True
        )
    )

    coluna1, coluna2, coluna3, coluna4 = (
        st.columns(4)
    )

    coluna1.metric(
        "Transportadores ativos",
        ativos
    )

    coluna2.metric(
        "Transportadores inativos",
        inativos
    )

    coluna3.metric(
        "Modalidades associadas",
        modalidades_ativas
    )

    ultimo_relatorio = (
        relatorios[0].get(
            "data_atualizacao",
            "Nenhum"
        )
        if relatorios
        else "Nenhum"
    )

    coluna4.metric(
        "Último relatório",
        ultimo_relatorio
    )

    st.divider()

    st.info(
        "Os cadastros são mantidos por CNPJ, permitindo "
        "associar várias modalidades à mesma empresa. "
        "Os relatórios já gerados permanecem preservados "
        "mesmo após futuras edições."
    )


def main():

    token, repo, branch, data_dir = (
        github_config()
    )

    if not token or not repo:

        st.error(
            "O aplicativo ainda não está configurado."
        )

        st.write(
            "Adicione nos Secrets do Streamlit: "
            "GITHUB_TOKEN e GITHUB_REPOSITORY."
        )

        st.stop()

    if "pagina" not in st.session_state:

        st.session_state[
            "pagina"
        ] = "Início"

    # -----------------------------------------------------------
    # Aplica qualquer navegação pedida por um botão na execução
    # anterior (via solicitar_pagina). Isso acontece ANTES do
    # st.sidebar.radio (mais abaixo) ser criado nesta execução,
    # então nunca conflita com o widget e nunca gera
    # StreamlitWidgetAlreadyInstantiatedError.
    # -----------------------------------------------------------

    if "pagina_solicitada" in st.session_state:

        st.session_state["pagina"] = (
            st.session_state.pop("pagina_solicitada")
        )

    if "editando_id" not in st.session_state:

        st.session_state[
            "editando_id"
        ] = None

    (
        transportadores,
        historico,
        relatorios,
        sha_t,
        sha_h,
        sha_r
    ) = carregar_dados_cache()

    st.sidebar.title(
        "Menu"
    )

    if st.sidebar.button(
        "🔄 Atualizar dados",
        use_container_width=True,
        help="Busca a versão mais recente dos dados diretamente no GitHub."
    ):

        carregar_dados_cache(
            forcar=True
        )

        st.rerun()

    st.sidebar.caption(
        f"{len(transportadores)} transportador(es) carregado(s)."
    )

    opcoes = [
        "Início",
        "Transportadores",
        "Novo Transportador",
        "Gerar Relatório",
        "Relatórios Anteriores"
    ]

    # O menu é ligado diretamente ao session_state através de
    # key="pagina" (sem usar "index="), que é a forma segura e
    # recomendada pelo Streamlit. A partir daqui, NADA neste script
    # pode voltar a escrever em st.session_state["pagina"] durante
    # esta mesma execução — qualquer navegação deve usar
    # solicitar_pagina(), acima.

    st.sidebar.radio(
        "Navegação",
        opcoes,
        key="pagina"
    )

    pagina_atual = st.session_state["pagina"]

    if pagina_atual == "Início":

        tela_inicio(
            transportadores,
            relatorios
        )

    elif pagina_atual == "Transportadores":

        tela_transportadores(
            transportadores,
            historico,
            sha_t,
            sha_h
        )

    elif pagina_atual == "Novo Transportador":

        tela_formulario(
            transportadores,
            historico,
            sha_t,
            sha_h
        )

    elif pagina_atual == "Gerar Relatório":

        tela_relatorio(
            transportadores,
            relatorios,
            sha_r
        )

    elif pagina_atual == "Relatórios Anteriores":

        tela_relatorios_anteriores(
            relatorios,
            sha_r
        )


if __name__ == "__main__":
    main()
